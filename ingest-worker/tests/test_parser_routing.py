"""Routing + timeout-registry contract for the Docling parsers (tickets #16/#17).

The un-gated tests are fast, need no emulators, and must NOT pull docling into
the test process (parser_for / parse_timeout_for are pure parent-side lookups).
The Docling-exec tests at the bottom are gated exactly like test_docling_convert.py
(DOCULYZE_RUN_DOCLING) because each conversion pays docling's ~10-30s import tax.
"""

import os
import subprocess
import sys

import pytest

from ingest import config
from ingest.parsers import (
    _CODE_LANG,
    _PARSER_BY_TYPE,
    PARSERS,
    CodeParser,
    HTMLParser,
    Parser,
    TextParser,
    XMLParser,
    parse_timeout_for,
    parser_for,
)
from ingest.precheck import CODE_FAMILY, DOCX_TYPE, TEXT_FAMILY


# --- routing (un-gated: pure lookups) ----------------------------------------

def test_html_routes_to_docling_html_parser():
    # #17: text/html leaves the bare-text family for the Docling html parser.
    assert parser_for("text/html") == "html"


def test_xml_routes_to_docling_xml_parser():
    # #17: application/xml leaves the code-fence family for the Docling xml parser.
    assert parser_for("application/xml") == "xml"


def test_docx_routes_to_docling_docx_parser():
    assert parser_for(DOCX_TYPE) == "docx"


def test_pdf_routes_to_pdf_parser():
    assert parser_for("application/pdf") == "pdf"


def test_plain_and_markdown_route_to_text():
    assert parser_for("text/plain") == "text"
    assert parser_for("text/markdown") == "text"


def test_code_family_except_xml_routes_to_fence_lang():
    # Every code type but application/xml still dispatches to its fence-lang parser.
    for ct in CODE_FAMILY:
        if ct == "application/xml":
            continue
        assert parser_for(ct) == _CODE_LANG[ct]


def test_every_signed_type_resolves_and_is_registered():
    # Routing table and PARSERS registry can't drift: every type the allowlist
    # can sign resolves without KeyError, and every resolved name has a parser.
    signed = TEXT_FAMILY | CODE_FAMILY | {"application/pdf", DOCX_TYPE}
    for ct in signed:
        name = parser_for(ct)  # KeyError here = routing/allowlist drift
        assert name in PARSERS, f"{ct} -> {name} missing from PARSERS"


def test_every_routing_value_has_a_registered_parser():
    # Every value the routing table can produce must be a key in the registry.
    for name in set(_PARSER_BY_TYPE.values()):
        assert name in PARSERS, f"{name} missing from PARSERS"


# --- registry contract (un-gated: construction only, no parse execution) ------

def test_xml_registered_as_xmlparser_not_codeparser():
    # Collision regression (#17): the CodeParser spread must NOT overwrite the
    # explicit XMLParser — "xml" is a Docling parser, not a fence lang.
    assert isinstance(PARSERS["xml"], XMLParser)
    assert not isinstance(PARSERS["xml"], CodeParser)


def test_html_registered_as_htmlparser():
    assert isinstance(PARSERS["html"], HTMLParser)


def test_every_parser_is_a_parser_instance_with_the_abc_contract():
    # Every registry value implements the Parser ABC: a callable `.parse` and a
    # boolean `docling` flag (the flag TextParser used to be missing, #14).
    for name, parser in PARSERS.items():
        assert isinstance(parser, Parser), name
        assert callable(parser.parse), name
        assert isinstance(parser.docling, bool), name


# --- timeout split (un-gated: config getters only) ---------------------------

def test_docling_parsers_carry_the_docling_budget():
    for name in ("docx", "html", "xml"):
        assert parse_timeout_for(name) == config.docling_timeout_s()


def test_non_docling_parsers_carry_the_parse_budget():
    fence_langs = {lang for ct, lang in _CODE_LANG.items() if ct != "application/xml"}
    for name in {"pdf", "text"} | fence_langs:
        assert parse_timeout_for(name) == config.parse_timeout_s()


def test_parse_timeout_for_resolves_for_every_parser_name():
    # Catches the TextParser bug (#14): a parser missing the `docling` attribute
    # made parse_timeout_for raise AttributeError instead of returning a budget.
    for name in PARSERS:
        budget = parse_timeout_for(name)
        assert isinstance(budget, float), name


# --- pure-seam parse behavior (un-gated: no docling, no subprocess) -----------

def test_text_parser_passes_bytes_through_as_utf8():
    assert PARSERS["text"].parse(b"hello") == "hello"


def test_code_parser_fences_with_language_hint():
    md = PARSERS["python"].parse(b"print('x')")
    assert md == "```python\nprint('x')\n```"


def test_code_fence_grows_past_embedded_backticks():
    # A payload containing a ``` run must get a longer opening/closing fence so it
    # can't break out (preserve the user's fence-length logic — test behavior).
    md = PARSERS["python"].parse(b"a\n```\nb")
    assert md.startswith("````python\n")  # 4 backticks > the 3 inside
    assert md.endswith("\n````")


# --- import-safety (un-gated: docling must NOT load on parsers import) --------

def test_module_import_does_not_pull_in_docling():
    # Importing ingest.parsers in a fresh interpreter must not import docling —
    # all docling imports live inside the parser functions (sandboxed-child pattern).
    # Constructing every PARSERS instance happens at module import, so this same
    # check also proves instance construction pulls in no heavy dep.
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    src = os.path.join(root, "src")
    code = (
        "import sys; import ingest.parsers; "
        "heavy = ('docling', 'pymupdf', 'docx', 'torch', 'onnxruntime'); "
        "leaked = sorted(m for m in sys.modules "
        "if any(m == h or m.startswith(h + '.') for h in heavy)); "
        "assert not leaked, leaked"
    )
    env = dict(os.environ, PYTHONPATH=src)
    proc = subprocess.run(
        [sys.executable, "-c", code], env=env, capture_output=True, text=True
    )
    assert proc.returncode == 0, proc.stderr


# --- Docling-exec at the pure PARSERS[name] seam (GATED) ----------------------

_GATE = pytest.mark.skipif(
    not os.environ.get("DOCULYZE_RUN_DOCLING"),
    reason="set DOCULYZE_RUN_DOCLING=1 to run Docling-exec tests (imports docling)",
)


@_GATE
def test_docx_parser_preserves_tables_and_headings():
    # #16 acceptance / the #11 table-drop bug class: a docx table must survive as
    # a Markdown pipe-table, and a heading as an ATX heading.
    import io

    from docx import Document

    buf = io.BytesIO()
    doc = Document()
    doc.add_heading("Report", level=1)
    doc.add_paragraph("Body paragraph text.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "AlphaCell"
    table.rows[0].cells[1].text = "BetaCell"
    table.rows[1].cells[0].text = "GammaCell"
    table.rows[1].cells[1].text = "DeltaCell"
    doc.save(buf)

    md = PARSERS["docx"].parse(buf.getvalue())

    assert "# Report" in md
    # A pipe-table row carrying the header cells separated by a pipe.
    assert any(
        "AlphaCell" in line and "BetaCell" in line and "|" in line
        for line in md.splitlines()
    ), md
    assert "GammaCell" in md and "DeltaCell" in md


@_GATE
def test_html_parser_extracts_heading_table_and_text():
    html = (
        b"<html><body>"
        b"<h1>Heading One</h1>"
        b"<p>Paragraph body text.</p>"
        b"<table><tr><td>CellOne</td><td>CellTwo</td></tr></table>"
        b"</body></html>"
    )
    md = PARSERS["html"].parse(html)

    assert "Heading One" in md
    assert "#" in md  # a heading marker survived
    assert "CellOne" in md and "CellTwo" in md
    # No raw input tags leaked into the Markdown payload.
    for tag in ("<h1>", "<p>", "<table>"):
        assert tag not in md, md


@_GATE
def test_xml_parser_surfaces_element_text_without_fence():
    # Generic XML through Docling's HTML backend (see _parse_xml rationale): the
    # element text content survives as Markdown, no code fence, no raw tags.
    # Tag names deliberately avoid HTML-semantic names (<title>, <body>, <head>):
    # Docling routes XML through BeautifulSoup's HTML parser, which treats those
    # as document metadata and drops their text. Neutral element names surface as
    # body content — this input reproduces the brief's empirically-verified output.
    xml = (
        b"<?xml version='1.0'?>"
        b"<report><section>Q1 Results</section>"
        b"<summary>Revenue grew 12 percent year over year.</summary>"
        b"<metrics><metric>Total</metric>"
        b"<amount>88.3</amount></metrics></report>"
    )
    md = PARSERS["xml"].parse(xml)

    assert "Q1 Results" in md
    assert "Revenue grew 12 percent year over year." in md
    assert "88.3" in md
    assert "```" not in md  # not code-fenced (that was the pre-#17 behavior)
    for tag in ("<report>", "<section>", "<metric>", "<amount>"):
        assert tag not in md, md
