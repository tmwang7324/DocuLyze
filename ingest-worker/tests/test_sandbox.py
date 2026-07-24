"""Sandboxed parser tests (guide Phase 6): a parser crash or garbage input is a
raised ParseError in the parent — never a worker-process death or a hang.

No emulators needed; fixtures are generated with the same libraries the parsers
use to WRITE files (writing is independent of the extraction under test).
"""

import io

import pytest

from ingest.parsers import parser_for
from ingest.sandbox import ParseError, run_parser_sandboxed


def pdf_fixture(html: str = "<p>plain body text</p>") -> bytes:
    # pypdf is gone (issue #11) — write the fixture with pymupdf, the same engine
    # pymupdf4llm reads. insert_htmlbox lets us produce a visibly-styled heading.
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page(width=400, height=400)
    page.insert_htmlbox(pymupdf.Rect(20, 20, 380, 260), html)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def docx_fixture(text: str) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_parser_dispatch_covers_the_allowlist():
    assert parser_for("application/pdf") == "pdf"
    assert (
        parser_for("application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        == "docx"
    )
    assert parser_for("text/plain") == "text"
    assert parser_for("application/json") == "json"  # code types are fenced
    with pytest.raises(KeyError):
        parser_for("application/msword")  # legacy .doc dropped in issue #11
    with pytest.raises(KeyError):
        parser_for("image/png")


def test_good_pdf_parses_to_markdown():
    # Visibly-styled heading + body. pymupdf4llm 1.28 emits "# **Big Heading**"
    # for an <h1> (verified empirically), so we can assert the `#` marker survives.
    md = run_parser_sandboxed(
        "pdf", pdf_fixture("<h1>Big Heading</h1><p>Some body paragraph text.</p>")
    )
    assert isinstance(md, str)
    assert "Big Heading" in md
    assert "Some body paragraph text." in md
    assert "#" in md  # heading marker preserved by the Markdown normalizer


def test_corrupt_pdf_raises_parse_error_not_hang():
    # Valid magic then garbage: passes precheck, must fail in the sandbox.
    with pytest.raises(ParseError):
        run_parser_sandboxed("pdf", b"%PDF-1.7" + bytes(range(256)) * 16)


def test_docx_parses_paragraph_text():
    assert "hello docx" in run_parser_sandboxed("docx", docx_fixture("hello docx"))


def test_docx_heading_becomes_markdown_heading():
    from docx import Document

    doc = Document()
    doc.add_heading("Section", level=2)
    buf = io.BytesIO()
    doc.save(buf)
    assert "## Section" in run_parser_sandboxed("docx", buf.getvalue())


def test_docx_table_survives_as_pipe_table():
    # Regression for the previously-dropped-table bug (#11): paragraphs-only
    # parsing silently discarded tables. Since #16 Docling owns docx and emits a
    # GitHub pipe table (with padded cells, e.g. "| h1   | h2   |"), in document
    # order between the intro and outro paragraphs. Assertions normalize the
    # per-cell padding so they test survival + order, not Docling's spacing.
    from docx import Document

    doc = Document()
    doc.add_paragraph("intro text")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "h1"
    t.cell(0, 1).text = "h2"
    t.cell(1, 0).text = "a"
    t.cell(1, 1).text = "b"
    doc.add_paragraph("outro text")
    buf = io.BytesIO()
    doc.save(buf)

    md = run_parser_sandboxed("docx", buf.getvalue())

    def _pipe_rows(text):
        # Collapse intra-cell whitespace so "| h1   | h2   |" -> "| h1 | h2 |".
        rows = []
        for line in text.splitlines():
            s = line.strip()
            if s.startswith("|") and s.endswith("|"):
                cells = [c.strip() for c in s.strip("|").split("|")]
                rows.append("| " + " | ".join(cells) + " |")
        return rows

    def _is_separator(row):
        cells = [c.strip() for c in row.strip("|").split("|")]
        return all(c and set(c) == {"-"} for c in cells)

    rows = _pipe_rows(md)
    assert "| h1 | h2 |" in rows
    assert any(_is_separator(r) for r in rows)  # the header/body separator row
    assert "| a | b |" in rows
    # document order preserved: intro before the table before outro
    assert md.index("intro text") < md.index("h1") < md.index("outro text")


def test_text_passthrough():
    assert run_parser_sandboxed("text", "héllo\n".encode("utf-8")) == "héllo\n"


def test_code_is_fenced_with_language_hint():
    src = '{"a": 1}\n'
    md = run_parser_sandboxed("json", src.encode("utf-8"))
    assert md == f"```json\n{src}\n```"


def test_code_fence_grows_past_embedded_backticks():
    # A JS template literal containing a ``` run must not break out of the fence:
    # the opening fence has to be longer than the longest run inside.
    src = "const s = `\n```\n`;\n"
    md = run_parser_sandboxed("javascript", src.encode("utf-8"))
    assert md.startswith("````javascript\n")  # 4 backticks > the 3 inside
    assert md.endswith("\n````")
