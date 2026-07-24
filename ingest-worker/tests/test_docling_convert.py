"""Docling groundwork (ticket #15): config knobs, import-safety, offline convert.

The un-gated tests must stay fast and NOT pull docling into the parent process —
importing `ingest.docling_convert` proves nothing loaded docling (all docling
imports live inside functions, the sandboxed-child pattern). The actual Docling
conversion is expensive (model/backends import ~10-30s), so the exec test is
gated behind DOCULYZE_RUN_DOCLING and runs offline.
"""

import os
import subprocess
import sys

import pytest

from ingest import config


# --- config knobs (un-gated: pure env reads) ---------------------------------

def test_docling_timeout_default():
    assert config.docling_timeout_s() == 60.0


def test_docling_timeout_env_override(monkeypatch):
    monkeypatch.setenv("DOCLING_TIMEOUT_S", "12.5")
    assert config.docling_timeout_s() == 12.5


def test_docling_artifacts_path_default_none(monkeypatch):
    # Unset -> None means "use Docling's own cache"; the container sets the env.
    monkeypatch.delenv("DOCLING_ARTIFACTS_PATH", raising=False)
    assert config.docling_artifacts_path() is None


def test_docling_artifacts_path_env_override(monkeypatch):
    monkeypatch.setenv("DOCLING_ARTIFACTS_PATH", "/opt/docling-models")
    assert config.docling_artifacts_path() == "/opt/docling-models"


# --- import-safety (un-gated: docling must NOT load on module import) ---------

def test_module_import_does_not_pull_in_docling():
    # Importing the converter module in a fresh interpreter must not import
    # docling (heavy). All docling imports live inside functions.
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    src = os.path.join(root, "src")
    code = (
        "import sys; import ingest.docling_convert; "
        "assert 'docling' not in sys.modules, "
        "sorted(m for m in sys.modules if m.startswith('docling'))"
    )
    env = dict(os.environ, PYTHONPATH=src)
    proc = subprocess.run(
        [sys.executable, "-c", code], env=env, capture_output=True, text=True
    )
    assert proc.returncode == 0, proc.stderr


# --- offline conversion (GATED: actually runs Docling) -----------------------

_GATE = pytest.mark.skipif(
    not os.environ.get("DOCULYZE_RUN_DOCLING"),
    reason="set DOCULYZE_RUN_DOCLING=1 to run Docling-exec tests (imports docling)",
)


@_GATE
def test_convert_html_stream_to_markdown_offline():
    from docling.datamodel.base_models import InputFormat

    from ingest.docling_convert import convert_to_markdown

    html = b"<html><body><h1>Title</h1><p>Hello world</p></body></html>"
    md = convert_to_markdown(html, filename="doc.html", input_format=InputFormat.HTML)
    assert "Hello world" in md
    assert "# Title" in md
    # No raw HTML tags survived the conversion.
    assert "<h1>" not in md and "<p>" not in md


@_GATE
def test_convert_docx_stream_to_markdown_offline():
    import io

    from docx import Document

    from docling.datamodel.base_models import InputFormat

    from ingest.docling_convert import convert_to_markdown

    buf = io.BytesIO()
    doc = Document()
    doc.add_heading("Report", level=1)
    doc.add_paragraph("Body paragraph text.")
    doc.save(buf)

    md = convert_to_markdown(
        buf.getvalue(), filename="doc.docx", input_format=InputFormat.DOCX
    )
    assert "Body paragraph text." in md
    assert "# Report" in md
